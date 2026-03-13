import importlib.util
from io import BytesIO
import unittest
import zipfile

from knowledge_base import extract_text_from_upload


HAS_DOCX = importlib.util.find_spec("docx") is not None


def build_minimal_docx_document_xml() -> str:
    return (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<w:document xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
        "<w:body>"
        "<w:p><w:r><w:t>회사 소개</w:t></w:r></w:p>"
        "<w:tbl>"
        "<w:tr>"
        "<w:tc><w:p><w:r><w:t>부서</w:t></w:r></w:p></w:tc>"
        "<w:tc><w:p><w:r><w:t>플랫폼</w:t></w:r></w:p></w:tc>"
        "</w:tr>"
        "<w:tr>"
        "<w:tc><w:p><w:r><w:t>이름</w:t></w:r></w:p></w:tc>"
        "<w:tc><w:p><w:r><w:t>김민재</w:t></w:r></w:p></w:tc>"
        "</w:tr>"
        "</w:tbl>"
        "</w:body>"
        "</w:document>"
    )


def build_minimal_docx_payload() -> bytes:
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, mode="w") as archive:
        archive.writestr("word/document.xml", build_minimal_docx_document_xml())
    return buffer.getvalue()


class KnowledgeUploadExtractionTests(unittest.TestCase):
    def test_extract_text_from_docx_archive_without_python_docx_dependency(self) -> None:
        payload = build_minimal_docx_payload()

        text = extract_text_from_upload("fallback.docx", payload)

        self.assertIn("회사 소개", text)
        self.assertIn("부서 | 플랫폼", text)
        self.assertIn("이름 | 김민재", text)

    @unittest.skipUnless(HAS_DOCX, "python-docx not installed")
    def test_extract_text_from_docx_includes_table_rows(self) -> None:
        from docx import Document  # type: ignore

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "부서"
        table.cell(0, 1).text = "플랫폼"
        table.cell(1, 0).text = "이름"
        table.cell(1, 1).text = "김민재"

        buffer = BytesIO()
        doc.save(buffer)

        text = extract_text_from_upload("onboarding.docx", buffer.getvalue())

        self.assertIn("부서 | 플랫폼", text)
        self.assertIn("이름 | 김민재", text)

    @unittest.skipUnless(HAS_DOCX, "python-docx not installed")
    def test_extract_text_from_docx_keeps_paragraphs_and_tables(self) -> None:
        from docx import Document  # type: ignore

        doc = Document()
        doc.add_paragraph("회사 소개")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "제품"
        table.cell(0, 1).text = "오케스트레이션"

        buffer = BytesIO()
        doc.save(buffer)

        text = extract_text_from_upload("company.docx", buffer.getvalue())

        self.assertIn("회사 소개", text)
        self.assertIn("제품 | 오케스트레이션", text)


if __name__ == "__main__":
    unittest.main()
